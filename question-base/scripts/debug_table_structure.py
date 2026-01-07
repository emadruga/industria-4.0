#!/usr/bin/env python3
"""Debug script to examine table structure in DOCX files."""

from docx import Document
from pathlib import Path

FILE_PATH = "/Users/emadruga/proj/industria-4.0/question-base/docs_by_author/EwertonMadruga/Estrutura e Gestão - Ewerton/20251105_checklist_gestao_agil.docx"

doc = Document(FILE_PATH)

print(f"Total tables: {len(doc.tables)}\n")

for idx, table in enumerate(doc.tables):
    print(f"{'='*80}")
    print(f"TABLE {idx}")
    print(f"{'='*80}")
    print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")

    # Print first few rows
    for row_idx, row in enumerate(table.rows[:5]):
        print(f"\n  Row {row_idx}:")
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text.strip()[:100]
            if text:
                print(f"    Cell {cell_idx}: {text}")

    # Check if contains evidence markers
    full_text = ""
    for row in table.rows:
        for cell in row.cells:
            full_text += cell.text + "\n"

    has_evidence = any(marker in full_text.lower() for marker in [
        'possíveis fontes de evidências',
        'fontes de evidência',
        'evidência',
        'evidencia'
    ])

    if has_evidence:
        print(f"\n  ✓ Contains evidence markers")
        print(f"\n  FULL TABLE TEXT (first 500 chars):")
        print(f"  {full_text[:500]}")

    print()
