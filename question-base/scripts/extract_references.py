#!/usr/bin/env python3
"""
References Extraction Module
Extracts academic/professional references from Industry 4.0 DOCX documents.

Author: Industry 4.0 Team
Version: 1.0.0
"""

import re
from typing import List, Dict, Optional
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


class ReferencesExtractor:
    """Extracts bibliographic references from DOCX documents."""

    def __init__(self):
        self.reference_markers = [
            'referências', 'referencias', 'references',
            'citação', 'citacao', 'citation',
            'bibliografia', 'bibliography'
        ]

    def extract_from_document(self, doc: Document) -> List[Dict[str, str]]:
        """
        Extract references from entire document.

        Args:
            doc: python-docx Document object

        Returns:
            List of dicts with 'citation' and optionally 'url' keys
        """
        references = []

        # Try extracting from tables first
        table_refs = self._extract_from_tables(doc)
        if table_refs:
            references.extend(table_refs)

        # Also try extracting from paragraphs (common format)
        para_refs = self._extract_from_paragraphs(doc)
        if para_refs:
            references.extend(para_refs)

        # Remove duplicates based on citation text
        references = self._deduplicate_references(references)

        return references

    def _extract_from_tables(self, doc: Document) -> List[Dict[str, str]]:
        """Extract references from tables in the document."""
        references = []

        for table in doc.tables:
            if self._is_references_table(table):
                # Skip header row
                for row in table.rows[1:]:
                    if len(row.cells) < 1:
                        continue

                    # Get raw text before cleaning (to preserve newlines)
                    raw_text = row.cells[0].text

                    # Check if this is a multi-line cell with multiple citations
                    if '\n' in raw_text:
                        # Split by newlines and process each as separate citation
                        lines = raw_text.split('\n')
                        for line in lines:
                            line = self._clean_text(line)
                            if self._is_valid_citation(line):
                                ref = self._parse_citation(line)
                                if ref:
                                    references.append(ref)
                    else:
                        # Single citation in cell
                        citation_text = self._clean_text(raw_text)
                        if self._is_valid_citation(citation_text):
                            ref = self._parse_citation(citation_text)
                            if ref:
                                references.append(ref)

        return references

    def _extract_from_paragraphs(self, doc: Document) -> List[Dict[str, str]]:
        """Extract references from paragraphs (common in some templates)."""
        references = []
        in_references_section = False

        for para in doc.paragraphs:
            text = para.text.strip()

            # Check if we've entered the references section
            if not in_references_section:
                if self._is_references_header(text):
                    in_references_section = True
                    continue

            # If in references section, extract citations
            if in_references_section and text:
                # Stop if we hit a new major section
                if self._is_new_section_header(text):
                    break

                # Check if this looks like a citation
                if self._is_valid_citation(text):
                    ref = self._parse_citation(text)
                    if ref:
                        references.append(ref)

        return references

    def _is_references_table(self, table: Table) -> bool:
        """Check if table contains references."""
        if len(table.rows) < 1:
            return False

        first_row_text = ' '.join([
            cell.text.lower() for cell in table.rows[0].cells
        ])

        return any(marker in first_row_text for marker in self.reference_markers)

    def _is_references_header(self, text: str) -> bool:
        """Check if text is a references section header."""
        text_lower = text.lower()

        # Must be short (headers are typically brief)
        if len(text) > 50:
            return False

        return any(marker in text_lower for marker in self.reference_markers)

    def _is_new_section_header(self, text: str) -> bool:
        """Check if text indicates a new section (end of references)."""
        # If text is very short and doesn't look like a citation
        if len(text) < 50 and not '.' in text:
            # Common section headers
            headers = ['glossário', 'anexo', 'apêndice', 'appendix']
            return any(h in text.lower() for h in headers)
        return False

    def _is_valid_citation(self, text: str) -> bool:
        """Check if text looks like a valid citation."""
        if not text or len(text) < 20:
            return False

        # Filter out header-like text
        text_lower = text.lower()
        if text_lower in ['citação', 'citation', 'referências', 'references']:
            return False

        # Valid citations usually have certain patterns
        # - Author names (uppercase letters)
        # - Year in parentheses or standalone
        # - Periods or commas
        has_uppercase = bool(re.search(r'[A-Z]{2,}', text))
        has_year = bool(re.search(r'\b(19|20)\d{2}\b', text))
        has_punctuation = '.' in text or ',' in text

        return has_uppercase and (has_year or has_punctuation)

    def _parse_citation(self, text: str) -> Optional[Dict[str, str]]:
        """Parse a citation and extract URL if present."""
        if not text:
            return None

        citation = self._clean_text(text)

        # Extract URL if present
        url_match = re.search(r'https?://[^\s\)\>]+', citation)

        result = {'citation': citation}

        if url_match:
            url = url_match.group(0)
            # Clean URL (remove trailing punctuation)
            url = url.rstrip('.,;:>')
            result['url'] = url

        return result

    def _deduplicate_references(self, references: List[Dict[str, str]]) -> List[Dict[str, str]]:
        """Remove duplicate references based on citation text."""
        seen = set()
        unique_refs = []

        for ref in references:
            # Use first 100 chars of citation as key
            key = ref['citation'][:100]

            if key not in seen:
                seen.add(key)
                unique_refs.append(ref)

        return unique_refs

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove markdown-like asterisks
        text = re.sub(r'\*+', '', text)

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Strip leading/trailing whitespace
        text = text.strip()

        return text


def extract_references_from_document(doc: Document) -> List[Dict[str, str]]:
    """
    Convenience function to extract references from a document.

    Args:
        doc: python-docx Document object

    Returns:
        List of reference entries
    """
    extractor = ReferencesExtractor()
    return extractor.extract_from_document(doc)


if __name__ == '__main__':
    # Test the extractor
    import sys
    from pathlib import Path

    if len(sys.argv) > 1:
        docx_path = Path(sys.argv[1])
        if docx_path.exists():
            doc = Document(str(docx_path))
            extractor = ReferencesExtractor()
            references = extractor.extract_from_document(doc)

            print(f"Found {len(references)} references:")
            for i, ref in enumerate(references[:5], 1):  # Show first 5
                print(f"  {i}. {ref['citation'][:80]}...")
                if 'url' in ref:
                    print(f"     URL: {ref['url']}")
        else:
            print(f"File not found: {docx_path}")
    else:
        print("References Extractor Module - Ready")
        print("Usage: python extract_references.py <docx_file>")
        print("   Or: from extract_references import ReferencesExtractor")
